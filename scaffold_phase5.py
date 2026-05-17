import os

files = {
    "reporting/schemas.py": """from pydantic import BaseModel, Field
from typing import List, Dict, Any, Optional
from datetime import datetime

class ReportCitation(BaseModel):
    segment_id: str
    exact_quote: str
    agent_id: str
    theory_applied: str
    reliability_score: float

class ReportSection(BaseModel):
    title: str
    content: str
    citations: List[ReportCitation] = Field(default_factory=list)
    visualizations: List[str] = Field(default_factory=list)
    ambiguity_warnings: List[str] = Field(default_factory=list)

class FinalResearchReport(BaseModel):
    corpus_id: str
    title: str
    executive_summary: str
    sections: List[ReportSection]
    overall_reliability: float
    generated_at: datetime = Field(default_factory=datetime.utcnow)
""",

    "reporting/generator.py": """import json
from typing import List, Dict, Any
from reporting.schemas import FinalResearchReport, ReportSection, ReportCitation
from reporting.synthesis import CrossAgentSynthesizer
from reporting.visualizations import VisualizationEngine

class ReportEngine:
    \"\"\"Assembles the final academic report from raw agent states.\"\"\"
    
    def __init__(self):
        self.synthesizer = CrossAgentSynthesizer()
        self.viz_engine = VisualizationEngine()
        
    async def generate_report(self, corpus_id: str, agent_states: List[Dict[str, Any]]) -> FinalResearchReport:
        # Synthesize findings
        synthesis = await self.synthesizer.synthesize(agent_states)
        
        # Generate Visualizations
        viz_paths = self.viz_engine.generate_all(synthesis)
        
        # Build sections
        sections = [
            ReportSection(
                title="Executive Summary",
                content=synthesis["summary"],
                ambiguity_warnings=synthesis["ambiguities"]
            ),
            ReportSection(
                title="Pragmatic Drift Analysis",
                content=synthesis["pragmatic_drift_text"],
                citations=synthesis["citations"],
                visualizations=[viz_paths.get("heatmap")] if viz_paths.get("heatmap") else []
            )
        ]
        
        return FinalResearchReport(
            corpus_id=corpus_id,
            title=f"Linguistic Analysis Report: Corpus {corpus_id}",
            executive_summary=synthesis["summary"],
            sections=sections,
            overall_reliability=synthesis["average_reliability"]
        )
""",

    "reporting/export.py": """import os
import json
from reporting.schemas import FinalResearchReport
import docx

class ExportPipeline:
    \"\"\"Handles DOCX and PDF export formatting.\"\"\"
    
    def export_docx(self, report: FinalResearchReport, output_path: str):
        doc = docx.Document()
        doc.add_heading(report.title, 0)
        
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(report.executive_summary)
        
        for section in report.sections:
            doc.add_heading(section.title, level=1)
            if section.ambiguity_warnings:
                p = doc.add_paragraph("WARNING: ")
                p.runs[0].bold = True
                p.add_run(" | ".join(section.ambiguity_warnings))
            
            doc.add_paragraph(section.content)
            
            if section.citations:
                doc.add_heading("Citations & Evidence", level=2)
                for cit in section.citations:
                    doc.add_paragraph(f"[{cit.segment_id}] '{cit.exact_quote}' (Reliability: {cit.reliability_score})", style='Quote')
                    
        doc.save(output_path)
        return output_path

    def export_markdown(self, report: FinalResearchReport, output_path: str):
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(f"# {report.title}\\n\\n")
            f.write(f"## Executive Summary\\n{report.executive_summary}\\n\\n")
            
            for section in report.sections:
                f.write(f"## {section.title}\\n")
                if section.ambiguity_warnings:
                    f.write(f"> **AMBIGUITY WARNING:** {', '.join(section.ambiguity_warnings)}\\n\\n")
                f.write(f"{section.content}\\n\\n")
        return output_path
""",

    "reporting/visualizations.py": """import os
from typing import Dict, Any

class VisualizationEngine:
    \"\"\"Generates Plotly/Matplotlib visual assets.\"\"\"
    
    def __init__(self, output_dir: str = "/tmp/pmdd_viz"):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)
        
    def generate_all(self, synthesis_data: Dict[str, Any]) -> Dict[str, str]:
        # In production, use Plotly to save static images
        heatmap_path = os.path.join(self.output_dir, "drift_heatmap.png")
        network_path = os.path.join(self.output_dir, "collocation_network.png")
        
        # Stubbing the actual image generation to prevent missing dependencies in base env
        with open(heatmap_path, "w") as f: f.write("MOCK PNG")
        with open(network_path, "w") as f: f.write("MOCK PNG")
        
        return {
            "heatmap": heatmap_path,
            "network": network_path
        }
""",

    "reporting/synthesis.py": """from typing import List, Dict, Any
from reporting.schemas import ReportCitation

class CrossAgentSynthesizer:
    \"\"\"Merges quantitative and qualitative data, reconciling contradictions.\"\"\"
    
    async def synthesize(self, agent_states: List[Dict[str, Any]]) -> Dict[str, Any]:
        # Mocking the LLM synthesis logic
        summary = "The corpus exhibits significant pragmatic drift toward direct directives."
        ambiguities = ["Agent 2 and Agent 3 disagreed on the implicature of 'community' in segment 45A."]
        
        citations = [
            ReportCitation(
                segment_id="45A",
                exact_quote="We must build the community.",
                agent_id="Agent2",
                theory_applied="Speech Act Theory",
                reliability_score=0.92
            )
        ]
        
        return {
            "summary": summary,
            "pragmatic_drift_text": "Detailed theoretical analysis goes here...",
            "ambiguities": ambiguities,
            "citations": citations,
            "average_reliability": 0.88
        }
""",

    "tests/reporting/test_report.py": """import pytest
import os
from reporting.generator import ReportEngine
from reporting.export import ExportPipeline

@pytest.mark.asyncio
async def test_report_generation_and_export():
    engine = ReportEngine()
    states = [{"agent": "A2", "findings": []}]
    
    report = await engine.generate_report("corpus-123", states)
    assert report.overall_reliability > 0.0
    assert len(report.sections) == 2
    
    exporter = ExportPipeline()
    md_path = "test_report.md"
    docx_path = "test_report.docx"
    
    exporter.export_markdown(report, md_path)
    exporter.export_docx(report, docx_path)
    
    assert os.path.exists(md_path)
    assert os.path.exists(docx_path)
    
    # Cleanup
    os.remove(md_path)
    os.remove(docx_path)
"""
}

def scaffold_phase5():
    # Install docx if not present (to prevent export failure)
    os.system("uv pip install python-docx plotly")
    
    for path, content in files.items():
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
    print("Phase 5 Research Reporting scaffolding complete.")

if __name__ == "__main__":
    scaffold_phase5()
