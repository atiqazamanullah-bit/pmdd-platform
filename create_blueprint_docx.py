import docx
from docx.shared import Pt
import re
import os

def create_docx(md_file, docx_file):
    doc = docx.Document()
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        if line.startswith("# "):
            p = doc.add_heading(line[2:].replace('*', ''), level=0)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].replace('*', ''), level=1)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].replace('*', ''), level=2)
        elif line.startswith("#### "):
            p = doc.add_heading(line[5:].replace('*', ''), level=3)
        elif line.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
            # basic bold parsing for the start of the line
            if text.startswith("**") and "**:" in text:
                parts = text.split("**:", 1)
                run = p.add_run(parts[0][2:] + ":")
                run.bold = True
                p.add_run(parts[1])
            else:
                p.add_run(text.replace('**', ''))
        elif line.startswith("---"):
            p = doc.add_paragraph()
            p.add_run("_" * 50)
        else:
            p = doc.add_paragraph()
            p.add_run(line.replace('**', ''))
            
    doc.save(docx_file)
    print(f"Created {docx_file}")

if __name__ == '__main__':
    md_path = r'C:\Users\gc\.gemini\antigravity\brain\cb123520-707d-428e-8b59-5daa889d69cd\PMDD_Production_Infrastructure_Blueprint.md'
    create_docx(md_path, 'PMDD_Production_Infrastructure_Blueprint.docx')
